"""
Document processing module for handling PDF and other document types
"""
import os
import logging
from typing import List, Dict, Any
from pathlib import Path
import PyPDF2
import fitz  # PyMuPDF
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
from config import Config

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document loading, parsing, and chunking"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            length_function=len,
        )
    
    def load_pdf_pypdf2(self, file_path: str) -> str:
        """Load PDF using PyPDF2"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Error loading PDF with PyPDF2: {e}")
            return ""
    
    def load_pdf_pymupdf(self, file_path: str) -> str:
        """Load PDF using PyMuPDF (fitz) - better for complex PDFs"""
        try:
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
            return text
        except Exception as e:
            logger.error(f"Error loading PDF with PyMuPDF: {e}")
            return ""
    
    def load_docx(self, file_path: str) -> str:
        """Load DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error loading DOCX: {e}")
            return ""
    
    def load_txt(self, file_path: str) -> str:
        """Load text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error loading TXT: {e}")
            return ""
    
    def load_document(self, file_path: str) -> str:
        """Load document based on file extension"""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            # Try PyMuPDF first, fallback to PyPDF2
            text = self.load_pdf_pymupdf(str(file_path))
            if not text.strip():
                text = self.load_pdf_pypdf2(str(file_path))
            return text
        elif extension == '.docx':
            return self.load_docx(str(file_path))
        elif extension == '.txt':
            return self.load_txt(str(file_path))
        else:
            logger.warning(f"Unsupported file format: {extension}")
            return ""
    
    def chunk_document(self, text: str, metadata: Dict[str, Any] = None) -> List[LangchainDocument]:
        """Split document into chunks"""
        if not text.strip():
            return []
        
        # Create metadata
        if metadata is None:
            metadata = {}
        
        # Create langchain document
        doc = LangchainDocument(page_content=text, metadata=metadata)
        
        # Split into chunks
        chunks = self.text_splitter.split_documents([doc])
        
        # Add chunk index to metadata
        for i, chunk in enumerate(chunks):
            chunk.metadata['chunk_index'] = i
        
        logger.info(f"Created {len(chunks)} chunks from document")
        return chunks
    
    def process_document(self, file_path: str) -> List[LangchainDocument]:
        """Complete document processing pipeline"""
        logger.info(f"Processing document: {file_path}")
        
        # Load document
        text = self.load_document(file_path)
        if not text.strip():
            logger.warning(f"No text extracted from {file_path}")
            return []
        
        # Create metadata
        file_path = Path(file_path)
        metadata = {
            'source': str(file_path),
            'filename': file_path.name,
            'file_type': file_path.suffix,
            'file_size': file_path.stat().st_size
        }
        
        # Chunk document
        chunks = self.chunk_document(text, metadata)
        
        return chunks
    
    def process_directory(self, directory_path: str) -> List[LangchainDocument]:
        """Process all supported documents in a directory"""
        directory_path = Path(directory_path)
        all_chunks = []
        
        for file_path in directory_path.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in Config.SUPPORTED_FORMATS:
                chunks = self.process_document(str(file_path))
                all_chunks.extend(chunks)
        
        logger.info(f"Processed {len(all_chunks)} total chunks from directory")
        return all_chunks

